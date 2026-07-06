import streamlit as st
from openai import OpenAI
import io
import re
from docx import Document

# --- 1. KONFIGURASI UTAMA STREAMLIT ---
st.set_page_config(
    page_title="GLM-5.2 Shared Workspace",
    page_icon="🔮",
    layout="wide"
)

# --- 2. FUNGSI UNTUK MEMBUAT FILE WORD (.DOCX) DENGAN FORMAT BENAR ---
def buat_file_word(riwayat_pesan):
    doc = Document()
    doc.add_heading('Draf Hasil Kerja AI - GLM-5.2 Workspace', level=0)
    
    for msg in riwayat_pesan:
        if msg["role"] == "system":
            continue
            
        if msg["role"] == "user":
            doc.add_heading("Pertanyaan / Instruksi Anda:", level=2)
            if isinstance(msg["content"], str) and "Lanjutkan penjelasan" not in msg["content"]:
                doc.add_paragraph(msg["content"])
                    
        elif msg["role"] == "assistant":
            doc.add_heading("Jawaban AI:", level=2)
            
            # Memisahkan teks berdasarkan baris agar paragraf & judul tetap rapi
            paragraf_list = msg["content"].split('\n')
            
            for p_text in paragraf_list:
                if not p_text.strip():
                    continue
                
                # --- DETEKSI & PEMBERSIH KODE PAGAR (#) ---
                match_heading = re.match(r'^(#{1,6})\s+(.*)$', p_text.strip())
                if match_heading:
                    level_pagar = len(match_heading.group(1))
                    teks_judul = match_heading.group(2)
                    teks_judul_bersih = teks_judul.replace('**', '')
                    level_word = min(level_pagar, 3) 
                    doc.add_heading(teks_judul_bersih, level=level_word)
                    continue
                
                # --- LOGIKA PEMBERSIH FORMAT BINTANG (**) PADA PARAGRAF ---
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', p_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        clean_text = part.replace('**', '')
                        p.add_run(clean_text).bold = True
                    else:
                        p.add_run(part)
                        
            p_line = doc.add_paragraph()
            p_line.add_run("_" * 40).italic = True
            
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 3. PANEL CONTROL SIDEBAR ---
with st.sidebar:
    st.title("🔮 Kontrol AI")
    st.info("⚡ Status Server: Terhubung Otomatis (GLM-5.2 Active)")
    
    st.divider()
    st.markdown("### 📥 Ekspor Dokumen")
    if "messages" in st.session_state and len(st.session_state.messages) > 1:
        file_word = buat_file_word(st.session_state.messages)
        st.download_button(
            label="📥 Download Jadi Word (.docx)",
            data=file_word,
            file_name="Draf_LagosAi_GLM52.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    else:
        st.info("Mulai obrolan terlebih dahulu untuk mengunduh berkas Word.")
            
    st.divider()
    if st.button("🗑️ Reset & Bersihkan Semua Memori"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

# --- 4. PEMASANGAN API KEY & KONFIGURASI GLM-5.2 NVIDIA ---
BASE_URL = "https://integrate.api.nvidia.com/v1"
nvidia_api_key = "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"
MODEL_NAME = "z-ai/glm-5.2"

client = OpenAI(base_url=BASE_URL, api_key=nvidia_api_key)

# --- 5. MANAJEMEN MEMORI CHAT ---
if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "system", "content": "Anda adalah GLM-5.2, model bahasa besar flagship untuk penalaran panjang, agen kecerdasan buatan, dan koding dari Z.ai yang berjalan di atas infrastruktur NVIDIA NIM. Jawab dalam Bahasa Indonesia secara terstruktur, cerdas, mendalam, dan natural."}
    ]

# --- 6. TAMPILAN UTAMA INTERFASE CHAT ---
st.title("🔮 Lagos AI 8.1 ")
st.caption("Workspace ditenagai oleh model z-ai/glm-5.2 dengan arsitektur terupdate dan dengan data yang di update secara realtime.")

for message in st.session_state.messages:
    if message["role"] != "system":
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

if len(st.session_state.messages) > 1 and st.session_state.messages[-1]["role"] == "assistant":
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📝 Lanjutkan Tulisan Ini", use_container_width=True):
            st.session_state.messages.append({"role": "user", "content": "Lanjutkan penjelasan tulisan Anda sebelumnya secara mengalir tanpa terputus."})
            st.rerun()

# --- 7. PROSES INPUT & RESPONS CHAT ---
user_input = st.chat_input("Ketik perintah teks Anda di sini...")

if user_input:
    st.session_state.messages.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.markdown(user_input)
        
    with st.chat_message("assistant"):
        try:
            response_stream = client.chat.completions.create(
                model=MODEL_NAME,
                messages=st.session_state.messages,
                temperature=0.3,
                max_tokens=2048,
                stream=True
            )
            
            def teks_generator_glm():
                for chunk in response_stream:
                    if hasattr(chunk, 'choices') and len(chunk.choices) > 0:
                        delta = chunk.choices[0].delta
                        content = getattr(delta, 'content', '')
                        if content:
                            yield content

            full_response = st.write_stream(teks_generator_glm())
            st.session_state.messages.append({"role": "assistant", "content": full_response})
            st.rerun()
            
        except Exception as e:
            st.error(f"Gagal memproses teks. Detail: {e}")
